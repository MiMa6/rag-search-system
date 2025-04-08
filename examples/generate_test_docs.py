"""
Generate test documents for RAG pipeline testing.
"""

import os
from datetime import datetime, timedelta
from pathlib import Path

import docx
from docx.shared import Inches
from fpdf import FPDF


def create_txt_file(filepath: str, content: str):
    with open(filepath, "w") as f:
        f.write(content)


def create_docx_file(filepath: str, content: str, title: str):
    doc = docx.Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    doc.save(filepath)


def create_pdf_file(filepath: str, content: str, title: str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=title, ln=1, align="C")
    pdf.multi_cell(0, 10, txt=content)
    pdf.output(filepath)


def main():
    # Create test documents directory
    base_dir = Path("rag_pipeline/data/test_docs")
    base_dir.mkdir(parents=True, exist_ok=True)

    # Sample document contents with versions
    document_sets = [
        {
            "name": "Project_Overview",
            "versions": [
                {
                    "date": "2023-01-15",
                    "content": """
Project Overview - Cloud Migration
Status: In Progress
Last Updated: January 15, 2023

Key Components:
1. Database Migration
2. Application Refactoring
3. Security Implementation

Timeline: Q1 2023 - Q3 2023
Budget: $500,000
                    """,
                },
                {
                    "date": "2023-06-20",
                    "content": """
Project Overview - Cloud Migration
Status: In Progress
Last Updated: June 20, 2023

Key Components:
1. Database Migration (Completed)
2. Application Refactoring (In Progress)
3. Security Implementation (Pending)
4. Performance Optimization (Added)

Timeline: Q1 2023 - Q4 2023
Budget: $650,000 (Revised)
                    """,
                },
            ],
        },
        {
            "name": "Technical_Specification",
            "versions": [
                {
                    "date": "2022-12-10",
                    "content": """
Technical Specification v1.0
Database Architecture

- PostgreSQL 13
- Redis Cache
- Daily Backup Schedule
- Standard Monitoring
                    """,
                },
                {
                    "date": "2023-03-15",
                    "content": """
Technical Specification v2.0
Database Architecture

- PostgreSQL 14
- Redis Cache with Clustering
- Hourly Backup Schedule
- Enhanced Monitoring
- Disaster Recovery Plan
                    """,
                },
            ],
        },
    ]

    # Generate documents in different formats
    for doc_set in document_sets:
        for version in doc_set["versions"]:
            base_name = f"{doc_set['name']}_{version['date']}"

            # Create TXT version
            create_txt_file(base_dir / f"{base_name}.txt", version["content"])

            # Create DOCX version
            create_docx_file(
                base_dir / f"{base_name}.docx",
                version["content"],
                f"{doc_set['name']} - {version['date']}",
            )

            # Create PDF version
            create_pdf_file(
                base_dir / f"{base_name}.pdf",
                version["content"],
                f"{doc_set['name']} - {version['date']}",
            )


if __name__ == "__main__":
    main()
    print("Test documents generated successfully!")
